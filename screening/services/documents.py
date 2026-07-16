"""Resume text extraction.

PDF goes through pypdf rather than pdfplumber: it is pure Python and far
smaller once installed, which keeps the serverless bundle under Vercel's limit
and the cold start short. Resumes are text-based PDFs, so the layout analysis
pdfplumber adds buys nothing here.
"""

import io
import os

from pypdf import PdfReader

SUPPORTED_FORMATS = (".pdf", ".docx")


class DocumentError(Exception):
    """The file could not be read, or carries no extractable text."""


def extract_text(data: bytes, filename: str = "") -> str:
    """Extract text from a resume, dispatching on the file's actual bytes.

    The extension is only a hint — a mislabelled .pdf that is really a DOCX
    still reads correctly, because the magic number decides.
    """
    if not data:
        raise DocumentError("The uploaded file is empty.")

    if data.lstrip()[:5].startswith(b"%PDF-"):
        return _from_pdf(data)

    # DOCX is a zip; "PK" is the local file header.
    if data[:2] == b"PK":
        return _from_docx(data)

    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        raise DocumentError("That file claims to be a PDF but isn't one.")

    raise DocumentError(
        f"Unsupported file. Expected {' or '.join(SUPPORTED_FORMATS)}."
    )


def _from_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:  # pypdf raises a wide range of parse errors
        raise DocumentError(f"This PDF could not be read: {exc}") from exc

    return _require_text("\n".join(p for p in pages if p.strip()))


def _from_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentError(
            "DOCX support needs python-docx. Install it, or upload a PDF."
        ) from exc

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentError(f"This DOCX could not be read: {exc}") from exc

    parts = [p.text for p in document.paragraphs]

    # Plenty of resumes lay their whole body out in a table.
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)

    return _require_text("\n".join(p for p in parts if p.strip()))


def _require_text(text: str) -> str:
    if not text.strip():
        raise DocumentError(
            "No text could be extracted. This looks like a scanned or "
            "image-only file, which needs OCR."
        )
    return text


def extract_text_from_path(path: str) -> str:
    with open(path, "rb") as f:
        return extract_text(f.read(), path)
